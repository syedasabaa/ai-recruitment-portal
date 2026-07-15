import fitz#fitz is what gives us the ability to OPEN and READ PDF files programmatically
import docx
import os

#1.Function to extract text from pdf
def extract_text_from_pdf(file_path:str)->str:
    text=""
    doc=fitz.open(file_path)
    for page in doc:
        text+=page.get_text()
    doc.close()
    return text.strip()


#2.Function to extract text from DOCX
def extract_text_from_docx(file_path:str)->str:
    doc=docx.Document(file_path)#opens word document at the given path
    text="\n".join([paragraph.text for paragraph in doc.paragraphs])#join()-->takes the list of string and merges them into one single string
    #list comprehension--->
    #paragraph_text=[]
    # for paragraph in doc.paragraphs
    #     paragraph_text.append(paragraph.text)
    return text.strip()


#3.Function for the Smart Dispatcher
def extract_text(file_path:str)->str:
    file_extension=os.path.splitext(file_path)[1].lower()
    if file_extension == ".pdf":
       return extract_text_from_pdf(file_path)
    elif file_extension == ".docx":
       return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_extension}")
